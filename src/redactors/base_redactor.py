"""Base class for all redactor implementations"""
from abc import ABC, abstractmethod
from typing import Dict, List, Tuple
from dataclasses import dataclass
from docx import Document

@dataclass
class PIIEntity:
    """Represents a detected PII entity"""
    text: str
    type: str
    start: int
    end: int
    confidence: float = 1.0

class BaseRedactor(ABC):
    """Abstract base class for PII redactors"""
    
    def __init__(self, consistency_mode: bool = True):
        """
        Initialize redactor
        
        Args:
            consistency_mode: If True, maintain consistent replacements for same values
        """
        self.consistency_mode = consistency_mode
        self.detected_entities: List[PIIEntity] = []
        self.replacements: Dict[str, str] = {}
    
    @abstractmethod
    def detect_pii(self, text: str) -> List[PIIEntity]:
        """
        Detect PII in text
        
        Args:
            text: Input text to analyze
        
        Returns:
            List of detected PII entities
        """
        pass
    
    @abstractmethod
    def generate_replacement(self, entity: PIIEntity) -> str:
        """
        Generate fake replacement for PII entity
        
        Args:
            entity: PIIEntity to replace
        
        Returns:
            Fake replacement value
        """
        pass
    
    def redact_text(self, text: str) -> Tuple[str, Dict[str, str]]:
        """
        Redact PII in text
        
        Args:
            text: Input text
        
        Returns:
            Tuple of (redacted_text, replacements_dict)
        """
        entities = self.detect_pii(text)
        
        # Sort by start position (reverse order for replacement)
        entities = sorted(entities, key=lambda x: x.start, reverse=True)
        
        redacted_text = text
        replacements = {}
        
        for entity in entities:
            original = entity.text
            
            # Use cached replacement if consistency mode is on
            if self.consistency_mode and original in self.replacements:
                fake = self.replacements[original]
            else:
                fake = self.generate_replacement(entity)
                self.replacements[original] = fake
            
            replacements[original] = fake
            
            # Replace in text
            redacted_text = redacted_text[:entity.start] + fake + redacted_text[entity.end:]
        
        return redacted_text, replacements
    
    def redact_document(self, doc: Document) -> Tuple[Document, Dict[str, str]]:
        """
        Redact PII in entire document
        
        Args:
            doc: Document object
        
        Returns:
            Tuple of (redacted_document, replacements_dict)
        """
        all_replacements = {}
        
        # Track seen entity texts to avoid counting duplicates (for statistics)
        seen_entity_texts = set()
        
        # Collect all entities first before redacting
        all_entities_with_positions = []
        
        # Redact paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                entities = self.detect_pii(para.text)
                
                # Store with paragraph reference
                for entity in entities:
                    all_entities_with_positions.append({
                        'entity': entity,
                        'para_idx': para_idx,
                        'para': para,
                        'is_table': False
                    })
        
        # Redact tables
        # Track processed cells to avoid duplicate processing of merged cells
        processed_cells = set()
        
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Skip if this cell was already processed (merged cells appear multiple times)
                    cell_id = id(cell)
                    if cell_id in processed_cells:
                        continue
                    processed_cells.add(cell_id)
                    
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para.text.strip():
                            entities = self.detect_pii(para.text)
                            
                            for entity in entities:
                                all_entities_with_positions.append({
                                    'entity': entity,
                                    'para_idx': para_idx,
                                    'para': para,
                                    'is_table': True,
                                    'table_idx': table_idx,
                                    'row_idx': row_idx,
                                    'cell_idx': cell_idx
                                })
        
        # Merge adjacent person names (fix fragmentation)
        all_entities_with_positions = self._merge_adjacent_person_names(all_entities_with_positions)
        
        # Store deduplicated entities for statistics
        for item in all_entities_with_positions:
            entity = item['entity']
            if entity.text not in seen_entity_texts:
                self.detected_entities.append(entity)
                seen_entity_texts.add(entity.text)
        
        # Now perform replacements
        for item in all_entities_with_positions:
            entity = item['entity']
            para = item['para']
            original = entity.text
            
            if self.consistency_mode and original in self.replacements:
                fake = self.replacements[original]
            else:
                fake = self.generate_replacement(entity)
                self.replacements[original] = fake
            
            all_replacements[original] = fake
            
            # Replace in paragraph runs
            for run in para.runs:
                if original in run.text:
                    run.text = run.text.replace(original, fake)
        
        return doc, all_replacements
    
    def _merge_adjacent_person_names(self, entities_with_positions: List[dict]) -> List[dict]:
        """
        Merge adjacent person name entities that are part of the same full name
        
        Args:
            entities_with_positions: List of entity dictionaries with position info
        
        Returns:
            Merged list with adjacent person names combined
        """
        if not entities_with_positions:
            return entities_with_positions
        
        # Group by paragraph
        by_para = {}
        for item in entities_with_positions:
            key = (item['para_idx'], item['is_table'])
            if key not in by_para:
                by_para[key] = []
            by_para[key].append(item)
        
        merged_result = []
        
        for key, items in by_para.items():
            # Sort by start position within paragraph
            items = sorted(items, key=lambda x: x['entity'].start)
            
            i = 0
            while i < len(items):
                current = items[i]
                current_entity = current['entity']
                
                # Only merge PERSON entities
                if current_entity.type != 'person':
                    merged_result.append(current)
                    i += 1
                    continue
                
                # Look ahead for adjacent person entities
                adjacent_persons = [current]
                j = i + 1
                
                while j < len(items):
                    next_item = items[j]
                    next_entity = next_item['entity']
                    
                    # Check if next is also a person and close by (within 3 chars - space or comma)
                    if (next_entity.type == 'person' and 
                        (next_entity.start - current_entity.end) <= 3):
                        adjacent_persons.append(next_item)
                        current_entity = next_entity  # Update for next distance check
                        j += 1
                    else:
                        break
                
                # If we found adjacent person names, merge them
                if len(adjacent_persons) > 1:
                    # Get the full text span from first to last
                    first_entity = adjacent_persons[0]['entity']
                    last_entity = adjacent_persons[-1]['entity']
                    
                    # Get the paragraph text to extract full name
                    para = current['para']
                    full_name = para.text[first_entity.start:last_entity.end]
                    
                    # Create merged entity
                    from .base_redactor import PIIEntity
                    merged_entity = PIIEntity(
                        text=full_name,
                        type='person',
                        start=first_entity.start,
                        end=last_entity.end,
                        confidence=sum(p['entity'].confidence for p in adjacent_persons) / len(adjacent_persons)
                    )
                    
                    # Use the first item's position info
                    merged_item = adjacent_persons[0].copy()
                    merged_item['entity'] = merged_entity
                    merged_result.append(merged_item)
                    
                    i = j  # Skip all merged entities
                else:
                    # No merge needed
                    merged_result.append(current)
                    i += 1
        
        return merged_result
    
    def get_statistics(self) -> Dict[str, int]:
        """Get statistics about detected PII"""
        stats = {}
        for entity in self.detected_entities:
            stats[entity.type] = stats.get(entity.type, 0) + 1
        return stats
    
    def reset(self):
        """Reset the redactor state"""
        self.detected_entities.clear()
        self.replacements.clear()
