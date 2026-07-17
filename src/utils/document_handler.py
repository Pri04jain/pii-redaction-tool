"""Handle DOCX document reading and writing operations"""
from docx import Document
from docx.shared import RGBColor
from typing import List, Dict, Tuple
import re

class DocumentHandler:
    """Handle DOCX document operations"""
    
    @staticmethod
    def read_document(file_path: str) -> Document:
        """Read DOCX document"""
        try:
            return Document(file_path)
        except Exception as e:
            raise ValueError(f"Error reading document: {str(e)}")
    
    @staticmethod
    def extract_text(doc: Document) -> str:
        """Extract all text from document"""
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return "\n".join(full_text)
    
    @staticmethod
    def get_paragraphs_with_metadata(doc: Document) -> List[Dict]:
        """Get paragraphs with metadata (position, style, etc.)"""
        paragraphs = []
        
        for idx, para in enumerate(doc.paragraphs):
            paragraphs.append({
                "index": idx,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "runs": [{"text": run.text, "bold": run.bold, "italic": run.italic} 
                        for run in para.runs]
            })
        
        return paragraphs
    
    @staticmethod
    def replace_text_in_document(doc: Document, replacements: Dict[str, str]) -> Document:
        """
        Replace text in document while preserving formatting
        
        Args:
            doc: Document object
            replacements: Dictionary mapping original text to replacement text
        
        Returns:
            Modified document
        """
        # Sort replacements by length (longest first) to avoid partial replacements
        sorted_replacements = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                for original, fake in sorted_replacements:
                    if original in run.text:
                        run.text = run.text.replace(original, fake)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for original, fake in sorted_replacements:
                                if original in run.text:
                                    run.text = run.text.replace(original, fake)
        
        # Replace in headers/footers
        for section in doc.sections:
            # Header
            header = section.header
            for para in header.paragraphs:
                for run in para.runs:
                    for original, fake in sorted_replacements:
                        if original in run.text:
                            run.text = run.text.replace(original, fake)
            
            # Footer
            footer = section.footer
            for para in footer.paragraphs:
                for run in para.runs:
                    for original, fake in sorted_replacements:
                        if original in run.text:
                            run.text = run.text.replace(original, fake)
        
        return doc
    
    @staticmethod
    def highlight_replacements(doc: Document, replacements: Dict[str, str], 
                              color: RGBColor = RGBColor(255, 255, 0)) -> Document:
        """Highlight replaced text in document (for review purposes)"""
        for para in doc.paragraphs:
            for run in para.runs:
                for fake in replacements.values():
                    if fake in run.text:
                        run.font.highlight_color = color
        
        return doc
    
    @staticmethod
    def save_document(doc: Document, output_path: str):
        """Save document to file"""
        try:
            doc.save(output_path)
        except Exception as e:
            raise ValueError(f"Error saving document: {str(e)}")
    
    @staticmethod
    def get_document_stats(doc: Document) -> Dict[str, int]:
        """Get document statistics"""
        return {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
            "total_chars": len(DocumentHandler.extract_text(doc))
        }
    
    @staticmethod
    def split_document(input_path: str, output_dir: str, num_parts: int = 6) -> List[str]:
        """
        Split a large document into smaller parts
        
        Args:
            input_path: Path to input DOCX
            output_dir: Directory for output files
            num_parts: Number of parts to split into
        
        Returns:
            List of output file paths
        """
        import os
        from pathlib import Path
        
        doc = Document(input_path)
        total_paragraphs = len(doc.paragraphs)
        paras_per_part = total_paragraphs // num_parts
        
        output_paths = []
        
        for i in range(num_parts):
            new_doc = Document()
            
            start_idx = i * paras_per_part
            end_idx = start_idx + paras_per_part if i < num_parts - 1 else total_paragraphs
            
            # Copy paragraphs
            for para in doc.paragraphs[start_idx:end_idx]:
                new_para = new_doc.add_paragraph(para.text)
                # Try to preserve basic formatting
                try:
                    new_para.style = para.style
                except:
                    pass  # Skip if style can't be applied
            
            # Save part
            output_path = os.path.join(output_dir, f"part_{i+1}.docx")
            new_doc.save(output_path)
            output_paths.append(output_path)
        
        return output_paths
