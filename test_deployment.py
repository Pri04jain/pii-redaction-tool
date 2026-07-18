#!/usr/bin/env python
"""Test script to diagnose Railway deployment issues"""

import requests
import json
import sys
from pathlib import Path

def test_deployment(base_url):
    """Test Railway deployment"""
    
    print("=" * 60)
    print("PII Redaction Tool - Deployment Test")
    print("=" * 60)
    print(f"Testing: {base_url}\n")
    
    # Test 1: Health check
    print("Test 1: Health Check")
    print("-" * 60)
    try:
        response = requests.get(f"{base_url}/health", timeout=10)
        if response.status_code == 200:
            health_data = response.json()
            print("✅ Health check passed")
            print(json.dumps(health_data, indent=2))
            
            # Check dependencies
            deps = health_data.get('dependencies', {})
            if not deps.get('spacy_model_loaded'):
                print("\n⚠️  WARNING: spaCy model not loaded!")
                print("This will cause redaction to fail.")
                
            if not deps.get('presidio_available'):
                print("\n⚠️  WARNING: Presidio not available!")
                print("Presidio and Hybrid modes will not work.")
                
        else:
            print(f"❌ Health check failed: {response.status_code}")
            print(response.text)
            return False
    except Exception as e:
        print(f"❌ Health check error: {e}")
        return False
    
    print("\n")
    
    # Test 2: Homepage
    print("Test 2: Homepage")
    print("-" * 60)
    try:
        response = requests.get(base_url, timeout=10)
        if response.status_code == 200:
            print("✅ Homepage accessible")
        else:
            print(f"❌ Homepage error: {response.status_code}")
    except Exception as e:
        print(f"❌ Homepage error: {e}")
    
    print("\n")
    
    # Test 3: API endpoints
    print("Test 3: API Endpoints")
    print("-" * 60)
    try:
        response = requests.get(f"{base_url}/api/modes", timeout=10)
        if response.status_code == 200:
            modes = response.json()
            print("✅ API modes endpoint working")
            print(f"Available modes: {len(modes.get('modes', []))}")
        else:
            print(f"❌ API error: {response.status_code}")
    except Exception as e:
        print(f"❌ API error: {e}")
    
    print("\n")
    
    # Test 4: File upload (if test file exists)
    print("Test 4: File Upload")
    print("-" * 60)
    
    # Create a simple test document
    try:
        from docx import Document
        
        # Create test document
        doc = Document()
        doc.add_paragraph("Test Document")
        doc.add_paragraph("Name: John Smith")
        doc.add_paragraph("Email: john.smith@example.com")
        doc.add_paragraph("Phone: 555-123-4567")
        
        test_file = Path("test_upload.docx")
        doc.save(test_file)
        
        print("Test document created")
        
        # Test with regex mode (simplest)
        for mode in ['regex', 'presidio', 'hybrid']:
            print(f"\nTesting {mode} mode...")
            try:
                with open(test_file, 'rb') as f:
                    files = {'file': ('test.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                    data = {'mode': mode}
                    response = requests.post(f"{base_url}/upload", files=files, data=data, timeout=60)
                    
                    if response.status_code == 200:
                        result = response.json()
                        if result.get('success'):
                            print(f"  ✅ {mode} mode: {result.get('total_redactions', 0)} redactions")
                            print(f"     Stats: {result.get('stats', {})}")
                        else:
                            print(f"  ❌ {mode} mode failed: {result}")
                    else:
                        print(f"  ❌ {mode} mode error: {response.status_code}")
                        print(f"     Response: {response.text}")
            except Exception as e:
                print(f"  ❌ {mode} mode exception: {e}")
        
        # Clean up
        test_file.unlink(missing_ok=True)
        
    except ImportError:
        print("⚠️  python-docx not installed, skipping upload test")
        print("   Install with: pip install python-docx")
    except Exception as e:
        print(f"❌ Upload test error: {e}")
    
    print("\n")
    print("=" * 60)
    print("Test Complete")
    print("=" * 60)
    
    return True


if __name__ == "__main__":
    if len(sys.argv) > 1:
        url = sys.argv[1]
    else:
        # Prompt for URL
        url = input("Enter your Railway app URL (e.g., https://your-app.railway.app): ").strip()
    
    # Remove trailing slash
    url = url.rstrip('/')
    
    # Validate URL
    if not url.startswith('http'):
        url = 'https://' + url
    
    test_deployment(url)
